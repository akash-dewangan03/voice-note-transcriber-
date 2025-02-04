import os
import time
import streamlit as st
import requests
from docx import Document 
# from transformers import pipeline
from src.transcriber import AudioToText
from dotenv import load_dotenv

load_dotenv()
api_key = os.getenv("ASSEMBLYAI_API_KEY")

st.set_page_config(page_title="🔉 Audio to Text 💬", page_icon="🔉")

with st.sidebar:
    st.title("🔉 Audio to Text 💬 Converter")

uploaded_file = st.file_uploader('Choose an audio file', type=['wav', 'mp3'])

if uploaded_file is not None and api_key:
    st.audio(uploaded_file, format='audio/wav')

    if st.button('Transcribe'):
        with st.spinner('Transcribing...'):
            transcription = AudioToText.transcribe_audio(uploaded_file, api_key)
            st.text_area('Transcription', transcription, height=300)

            # Summarization Pipeline
            summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=0)
            summary = summarizer(transcription, max_length=150, min_length=50, do_sample=False)[0]["summary_text"]

            # Convert summary to bullet points
            bullet_summary = "\n".join(f"- {point.strip()}" for point in summary.split(". ") if point)

            st.subheader("🔹 Bullet-Point Summary")
            st.text_area("Summary", bullet_summary, height=200)

            # Function to create DOCX
            def create_docx(summary_text):
                doc = Document()
                doc.add_heading("Audio Summary", level=1)

                for point in summary_text.split("\n"):
                    doc.add_paragraph(point, style="ListBullet")

                doc_path = "summary.docx"
                doc.save(doc_path)
                return doc_path

            docx_file = create_docx(bullet_summary)

            # Download Summary as DOCX
            with open(docx_file, "rb") as file:
                st.download_button(
                    label="Download Summary as DOCX",
                    data=file,
                    file_name="summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            # Download Summary as TXT
            st.download_button(
                label="Download Summary as TXT",
                data=bullet_summary,
                file_name="summary.txt",
                mime="text/plain"
            )
